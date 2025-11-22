from docx import Document
from docx.shared import Pt, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_ebook(title, preface_content, chapters_content, conclusion_content):
    doc = Document()
    
    # Set A5 Size (148mm x 210mm)
    section = doc.sections[0]
    section.page_height = Mm(210)
    section.page_width = Mm(148)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    # ========== TYPOGRAPHY SETUP ==========
    # Normal text style
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    
    # Set paragraph format for normal style
    normal_para_format = normal_style.paragraph_format
    normal_para_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal_para_format.line_spacing = 1.15
    normal_para_format.space_after = Pt(6)
    normal_para_format.first_line_indent = Inches(0)  # NO indent - flush left for all
    normal_para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Explicitly LEFT
    
    # Heading styles
    heading_configs = {
        'Heading 1': {'size': Pt(16), 'bold': True, 'color': RGBColor(0, 51, 102)},
        'Heading 2': {'size': Pt(14), 'bold': True, 'color': RGBColor(0, 51, 102)},
        'Heading 3': {'size': Pt(12), 'bold': True, 'color': RGBColor(0, 51, 102)},
    }
    
    for style_name, config in heading_configs.items():
        try:
            heading_style = doc.styles[style_name]
            heading_font = heading_style.font
            heading_font.name = 'Calibri'
            heading_font.size = config['size']
            heading_font.bold = config['bold']
            heading_font.color.rgb = config['color']
            
            # Heading spacing and alignment
            heading_para = heading_style.paragraph_format
            heading_para.space_before = Pt(12)
            heading_para.space_after = Pt(6)
            
            # Center-align Heading 1 (chapter titles), others LEFT
            if style_name == 'Heading 1':
                heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except:
            pass
    
    # List Bullet style
    try:
        list_style = doc.styles['List Bullet']
        list_font = list_style.font
        list_font.name = 'Calibri'
        list_font.size = Pt(11)
        # Ensure left alignment for lists
        list_para = list_style.paragraph_format
        list_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except:
        pass
    
    # ========== TITLE PAGE ==========
    title_heading = doc.add_heading(title, 0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_heading.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_page_break()

    # ========== KATA PENGANTAR ==========
    _add_section_with_content(doc, "Kata Pengantar", preface_content)
    doc.add_page_break()

    # ========== DAFTAR ISI ==========
    _add_table_of_contents(doc, chapters_content)
    
    # ========== PAGE NUMBERING ==========
    _add_page_numbers(doc)
    
    # ========== CHAPTERS ==========
    for i, (chapter_title, content) in enumerate(chapters_content, 1):
        doc.add_page_break()
        
        # Chapter heading - CENTER ALIGNED
        chapter_heading = doc.add_heading(f"Bab {i}: {chapter_title}", level=1)
        chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # CENTER for chapter titles
        
        # Add spacing after chapter title
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)
        
        # Clean content before parsing (remove excessive blank lines)
        cleaned_content = _clean_excessive_spacing(content)
        
        # Parse and add chapter content
        _parse_and_add_content(doc, cleaned_content, is_first_chapter_para=True)

    # ========== PENUTUP ==========
    doc.add_page_break()
    _add_section_with_content(doc, "Penutup", conclusion_content)

    return doc

def _add_section_with_content(doc, section_title, content):
    """Add a section (Preface/Conclusion) with proper formatting"""
    heading = doc.add_heading(section_title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)
    
    # Clean duplicate section titles from content
    # AI sometimes writes "Kata Pengantar" or "Penutup" again in the content
    cleaned_content = content
    if section_title.lower() in content.lower()[:100]:  # Check first 100 chars
        # Remove lines that are just the section title (with or without quotes/formatting)
        lines = content.split('\n')
        filtered_lines = []
        for line in lines:
            clean_line = line.strip().replace('"', '').replace('*', '').replace(':', '').strip()
            if clean_line.lower() != section_title.lower():
                filtered_lines.append(line)
        cleaned_content = '\n'.join(filtered_lines)
    
    _parse_and_add_content(doc, cleaned_content, is_first_chapter_para=True)

def _add_table_of_contents(doc, chapters_content):
    """Add automatic Table of Contents"""
    doc.add_heading("Daftar Isi", level=1)
    
    # Add TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    # Add instruction
    instruction = doc.add_paragraph("(Klik kanan dan pilih 'Update Field' untuk memperbarui)")
    instruction.runs[0].font.size = Pt(9)
    instruction.runs[0].font.italic = True
    instruction.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # doc.add_page_break() - Removed to prevent double page break before Chapter 1

def _add_page_numbers(doc):
    """Add page numbers to footer"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = footer_para.add_run()
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def _clean_excessive_spacing(content):
    """Remove excessive blank lines and underscores from content"""
    # First, remove all underscores
    content = content.replace('_', '')
    
    lines = content.split('\n')
    cleaned_lines = []
    consecutive_blanks = 0
    
    for line in lines:
        if line.strip() == '':
            consecutive_blanks += 1
            # Allow maximum 1 consecutive blank line
            if consecutive_blanks <= 1:
                cleaned_lines.append(line)
        else:
            consecutive_blanks = 0
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def _parse_and_add_content(doc, content, is_first_chapter_para=False):
    """Parse Markdown content and add to document with professional formatting"""
    lines = content.split('\n')
    i = 0
    first_paragraph = is_first_chapter_para
    in_action_plan = False
    
    while i < len(lines):
        line = lines[i].strip()
        i += 1
        
        if not line:
            continue
        
        # Skip conversational filler
        lower_line = line.lower()
        if (lower_line.startswith("tentu") or lower_line.startswith("ini dia") or 
            lower_line.startswith("berikut") or lower_line.startswith("baik")):
            continue

        # ========== HEADINGS ==========
        if line.startswith('#'):
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1
            text = line[level:].strip()
            
            # Skip structural meta-headings AND duplicate section titles
            text_upper = text.upper()
            skip_headings = [
                "THE HOOK", "THE BODY", "GOLDEN QUOTE", "ACTION PLAN", 
                "A. THE HOOK", "B. THE BODY", "C. GOLDEN QUOTE", "D. ACTION PLAN",
                "HOOK", "BODY", "ISI", "PEMBUKAAN", "ISI UTAMA",
                "KATA PENGANTAR", "PENUTUP", "CONCLUSION", "PREFACE"
            ]
            if any(skip in text_upper for skip in skip_headings):
                continue
            
            # Clean asterisks from heading text
            text = text.replace('**', '').replace('*', '').strip()
            
            word_level = min(level + 1, 3)
            doc.add_heading(text, level=word_level)
            first_paragraph = True  # Reset for next paragraph
            continue
        
        # ========== ACTION PLAN DETECTION ==========
        if "LANGKAH NYATA" in line.upper() or "MISI HARI INI" in line.upper() or "RENCANA AKSI" in line.upper():
            # Add horizontal separator before action plan
            _add_horizontal_line(doc)
            
            # Action plan heading with shaded background
            action_heading = doc.add_paragraph()
            action_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = action_heading.add_run(line.replace('*', '').upper())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Add shading to paragraph
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E8F4F8')  # Light blue background
            action_heading._element.get_or_add_pPr().append(shading_elm)
            
            action_heading.paragraph_format.space_before = Pt(12)
            action_heading.paragraph_format.space_after = Pt(8)
            
            in_action_plan = True
            continue
        
        # ========== NUMBERED/BULLET LISTS ==========
        is_list_item = False
        list_content = line
        
        if len(line) > 2 and line[0].isdigit() and line[1] in ['.', ')']:
            is_list_item = True
            list_content = line.split('.', 1)[1].strip() if '.' in line else line.split(')', 1)[1].strip()
        elif line.startswith('- ') or line.startswith('• ') or line.startswith('* '):
            is_list_item = True
            list_content = line[2:].strip()
        
        if is_list_item:
            p = doc.add_paragraph(style='List Bullet')
            
            # Add checkbox for action plan items
            if in_action_plan:
                checkbox_run = p.add_run('☐ ')
                checkbox_run.font.size = Pt(12)
            
            _add_formatted_text(p, list_content)
            
            # Shading for action plan lists
            if in_action_plan:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F5F5F5')  # Light gray
                p._element.get_or_add_pPr().append(shading_elm)
            
            continue
        
        # ========== REGULAR PARAGRAPHS ==========
        p = doc.add_paragraph()
        
        # First paragraph: NO indent + Drop cap (if first of chapter)
        if first_paragraph and not in_action_plan:
            p.paragraph_format.first_line_indent = Inches(0)  # Already 0 by default now
            first_paragraph = False
        
        # Check if it's a Golden Quote (standalone bold text)
        if line.count('**') == 2 and len(line) < 200:
            # Golden Quote Box
            clean_text = line.replace('**', '').strip()
            
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            
            run = p.add_run(f'"{clean_text}"')
            run.italic = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Add light background
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F8FF')  # Alice blue
            p._element.get_or_add_pPr().append(shading_elm)
            
            # Add border
            pBdr = OxmlElement('w:pBdr')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), 'B0C4DE')
                pBdr.append(border)
            p._element.get_or_add_pPr().append(pBdr)
            
            continue
        
        # Regular paragraph with formatting
        _add_formatted_text(p, line)

def _add_horizontal_line(doc):
    """Add a horizontal line separator"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Thickness
    bottom.set(qn('w:color'), '0066CC')  # Blue color
    pBdr.append(bottom)
    pPr.append(pBdr)

def _add_formatted_text(paragraph, text):
    """Add text to paragraph with bold formatting support"""
    # Remove underscores first
    text = text.replace('_', '')
    text = text.strip()
    parts = text.split("**")
    
    for i, part in enumerate(parts):
        if not part:
            continue
        
        clean_part = part.replace('*', '').replace('"', '"').replace('"', '"')
        run = paragraph.add_run(clean_part)
        
        if i % 2 == 1:
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
